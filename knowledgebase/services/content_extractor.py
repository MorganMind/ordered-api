from typing import Dict, Any, List, Optional
import PyPDF2
import docx
from typing import List, Dict, Any
from common.logger.logger_service import get_logger
from datetime import datetime
import os
from .code_extractors.factory import CodeExtractorFactory
from .content_types import TextSegment, DocumentSection, DocumentMetadata, ExtractedContent

class ContentExtractor:
    logger = get_logger()

    # Class-level constants for supported types
    PDF_TYPES = {
        'extensions': ['.pdf', 'pdf'],
        'mime_types': ['application/pdf']
    }
    
    DOCX_TYPES = {
        'extensions': ['.docx', 'docx'],
        'mime_types': [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/docx'
        ]
    }
    
    TXT_TYPES = {
        'extensions': ['.txt', 'txt'],
        'mime_types': ['text/plain']
    }

    CODE_TYPES = {
        'extensions': [
            '.py', '.js', '.ts', '.jsx', '.tsx', '.java', 
            '.cpp', '.cs', '.rb', '.go', '.rs'
        ],
        'mime_types': [
            'text/x-python',
            'application/javascript',
            'text/javascript',
            'text/jsx',
            'text/tsx',
            'application/jsx',
            'application/tsx',
            'text/x-java-source',
            'text/x-c++src'
        ]
    }

    @classmethod
    async def extract_text(
        cls,
        file_path: str,
        file_type: str,
        extract_metadata: bool = True,
        generate_report: bool = False,
        report_path: Optional[str] = None
    ) -> ExtractedContent:
        """
        Extracts text and metadata from a file
        Returns standardized ExtractedContent object
        """
        try:
            file_type_lower = file_type.lower()
            
            # Extract content based on file type
            if file_type_lower in cls.PDF_TYPES['extensions'] or file_type_lower in cls.PDF_TYPES['mime_types']:
                extracted_content = await cls._extract_from_pdf(file_path)
            elif file_type_lower in cls.DOCX_TYPES['extensions'] or file_type_lower in cls.DOCX_TYPES['mime_types']:
                extracted_content = await cls._extract_from_docx(file_path)
            elif file_type_lower in cls.TXT_TYPES['extensions'] or file_type_lower in cls.TXT_TYPES['mime_types']:
                extracted_content = await cls._extract_from_txt(file_path)
            elif file_type_lower in cls.CODE_TYPES['extensions'] or file_type_lower in cls.CODE_TYPES['mime_types']:
                extracted_content = await cls._extract_from_code(file_path)
            else:
                supported_types = {
                    'PDF': cls.PDF_TYPES,
                    'DOCX': cls.DOCX_TYPES,
                    'TXT': cls.TXT_TYPES,
                    'CODE': cls.CODE_TYPES
                }
                raise ValueError(
                    f"Unsupported file type: {file_type}. "
                    f"Supported types: {supported_types}"
                )
        
            # Generate report if requested
            if generate_report:
                if report_path is None:
                    base_path = os.path.splitext(file_path)[0]
                    report_path = f"{base_path}_extraction_report.html"
                
                cls.generate_extraction_report(extracted_content, report_path)
                cls.logger.info(f"Generated extraction report at: {report_path}")
                
            return extracted_content

        except Exception as e:
            cls.logger.error("Text extraction failed", extra={
                "file_type": file_type,
                "error": str(e)
            })
            raise

    @classmethod
    async def _extract_from_pdf(cls, file_path: str) -> ExtractedContent:
        """Extract text and metadata from PDF"""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            sections: List[DocumentSection] = []
            full_text: List[str] = []
            total_chars = 0
            
            # Each page becomes a section
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if not text.strip():
                    continue
                
                # Split page text into paragraphs
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                
                section_segments: List[TextSegment] = []
                section_char_count = 0
                
                for para_idx, paragraph in enumerate(paragraphs):
                    segment = TextSegment(
                        text=paragraph,
                        segment_type='paragraph',
                        segment_index=para_idx,
                        char_count=len(paragraph),
                        metadata={
                            'page_number': page_num + 1,
                            'paragraph_number': para_idx + 1
                        }
                    )
                    section_segments.append(segment)
                    section_char_count += segment.char_count
                    full_text.append(paragraph)
                    total_chars += segment.char_count
                
                section = DocumentSection(
                    heading=f"Page {page_num + 1}",
                    segments=section_segments,
                    section_index=page_num,
                    char_count=section_char_count
                )
                sections.append(section)

            # Try to extract PDF metadata
            metadata = DocumentMetadata(
                total_chars=total_chars,
                sections=sections,
                document_type='pdf',
                creation_date=reader.metadata.get('/CreationDate', None),
                last_modified_date=reader.metadata.get('/ModDate', None),
                author=reader.metadata.get('/Author', None),
                title=reader.metadata.get('/Title', None),
                page_count=len(reader.pages),
                language=None  # Could be extracted from PDF language metadata if available
            )

            return ExtractedContent(
                text='\n\n'.join(full_text),
                metadata=metadata
            )

    @classmethod
    async def _extract_from_docx(cls, file_path: str) -> ExtractedContent:
        """Extract text and metadata from DOCX including sections, headers, and columns"""
        doc = docx.Document(file_path)
        sections: List[DocumentSection] = []
        full_text: List[str] = []
        total_chars = 0
        
        # Extract document properties
        core_props = doc.core_properties

        # Process each section in the document
        for section_idx, doc_section in enumerate(doc.sections):
            section_segments: List[TextSegment] = []
            section_text: List[str] = []
            segment_idx = 0

            # Extract headers
            for header in doc_section.header.paragraphs:
                if header.text.strip():
                    segment = TextSegment(
                        text=header.text.strip(),
                        segment_type='header',
                        segment_index=segment_idx,
                        char_count=len(header.text.strip()),
                        metadata={
                            'style': header.style.name if header.style else None,
                            'alignment': header.alignment
                        }
                    )
                    section_segments.append(segment)
                    section_text.append(header.text.strip())
                    total_chars += segment.char_count
                    segment_idx += 1

            # Extract main content - iterate through paragraphs and tables
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Determine if it's a heading
                    is_heading = paragraph.style and 'Heading' in paragraph.style.name
                    
                    # Check for column breaks or other special formatting in the paragraph's runs
                    runs_text = []
                    for run in paragraph.runs:
                        if run.text.strip():
                            runs_text.append(run.text.strip())
                    
                    # If we found individual runs, use them; otherwise use paragraph text
                    final_text = text if not runs_text else ' '.join(runs_text)
                    
                    segment = TextSegment(
                        text=final_text,
                        segment_type='heading' if is_heading else 'paragraph',
                        segment_index=segment_idx,
                        char_count=len(final_text),
                        metadata={
                            'style': paragraph.style.name if paragraph.style else None,
                            'alignment': paragraph.alignment,
                            'in_column': True  # Assume text is in columns unless proven otherwise
                        }
                    )
                    section_segments.append(segment)
                    section_text.append(final_text)
                    total_chars += segment.char_count
                    segment_idx += 1

            # Extract tables
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            segment = TextSegment(
                                text=cell.text.strip(),
                                segment_type='table_cell',
                                segment_index=segment_idx,
                                char_count=len(cell.text.strip()),
                                metadata={
                                    'row': row_idx,
                                    'column': cell_idx,
                                    'table_index': len(section_segments)
                                }
                            )
                            section_segments.append(segment)
                            section_text.append(cell.text.strip())
                            total_chars += segment.char_count
                            segment_idx += 1

            # Extract footers
            for footer in doc_section.footer.paragraphs:
                if footer.text.strip():
                    segment = TextSegment(
                        text=footer.text.strip(),
                        segment_type='footer',
                        segment_index=segment_idx,
                        char_count=len(footer.text.strip()),
                        metadata={
                            'style': footer.style.name if footer.style else None,
                            'alignment': footer.alignment
                        }
                    )
                    section_segments.append(segment)
                    section_text.append(footer.text.strip())
                    total_chars += segment.char_count
                    segment_idx += 1

            # Create section with all its segments
            section = DocumentSection(
                heading=f"Section {section_idx + 1}",
                segments=section_segments,
                section_index=section_idx,
                char_count=sum(len(text) for text in section_text)
            )
            sections.append(section)
            full_text.extend(section_text)

        metadata = DocumentMetadata(
            total_chars=total_chars,
            sections=sections,
            document_type='docx',
            creation_date=core_props.created if core_props else None,
            last_modified_date=core_props.modified if core_props else None,
            author=core_props.author if core_props else None,
            title=core_props.title if core_props else None,
            page_count=len(doc.sections),
            language=None  # Could be extracted from document language settings
        )

        return ExtractedContent(
            text='\n\n'.join(full_text),
            metadata=metadata
        )

    @classmethod
    async def _extract_from_txt(cls, file_path: str) -> ExtractedContent:
        """Extract text and metadata from TXT"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            
            # Split into paragraphs (double newlines)
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            # Create a single section for the entire text file
            segments: List[TextSegment] = []
            total_chars = 0
            
            for idx, paragraph in enumerate(paragraphs):
                segment = TextSegment(
                    text=paragraph,
                    segment_type='paragraph',
                    segment_index=idx,
                    char_count=len(paragraph),
                    metadata={
                        'paragraph_number': idx + 1,
                        'line_count': len(paragraph.split('\n'))
                    }
                )
                segments.append(segment)
                total_chars += segment.char_count

            section = DocumentSection(
                heading=None,
                segments=segments,
                section_index=0,
                char_count=total_chars
            )

            # Get file metadata
            file_stats = os.stat(file_path)
            metadata = DocumentMetadata(
                total_chars=total_chars,
                sections=[section],
                document_type='txt',
                creation_date=datetime.fromtimestamp(file_stats.st_ctime),
                last_modified_date=datetime.fromtimestamp(file_stats.st_mtime),
                author=None,
                title=os.path.basename(file_path),
                page_count=1,  # TXT files are considered single-page
                language=None  # Could be detected using langdetect or similar
            )

            return ExtractedContent(
                text=text,
                metadata=metadata
            )

    @classmethod
    async def _extract_from_code(cls, file_path: str) -> ExtractedContent:
        """Extract text and metadata from code files"""
        file_extension = os.path.splitext(file_path)[1]
        extractor = CodeExtractorFactory.get_extractor(file_extension)
        return await extractor.extract(file_path)

    @staticmethod
    def generate_extraction_report(extracted_content: ExtractedContent, output_path: str) -> str:
        """Generate an HTML report of the extracted content and metadata"""
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Content Extraction Report</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .metadata {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin: 10px 0; }}
                .section {{ border: 1px solid #ddd; margin: 10px 0; padding: 10px; border-radius: 5px; }}
                .segment {{ margin: 10px 0; padding: 10px; background: #fff; border: 1px solid #eee; }}
                .collapsible {{ 
                    background: #eee;
                    cursor: pointer;
                    padding: 18px;
                    width: 100%;
                    border: none;
                    text-align: left;
                    outline: none;
                    font-size: 15px;
                }}
                .active, .collapsible:hover {{ background-color: #ddd; }}
                .content {{ 
                    padding: 0 18px;
                    display: none;
                    overflow: hidden;
                    background-color: #f1f1f1;
                }}
                .metadata-item {{ margin: 5px 0; }}
                pre {{ white-space: pre-wrap; }}
            </style>
        </head>
        <body>
            <h1>Content Extraction Report</h1>
            
            <button class="collapsible">Document Metadata</button>
            <div class="content">
                <div class="metadata">
                    <div class="metadata-item"><b>Document Type:</b> {extracted_content.metadata.document_type}</div>
                    <div class="metadata-item"><b>Total Characters:</b> {extracted_content.metadata.total_chars}</div>
                    <div class="metadata-item"><b>Page Count:</b> {extracted_content.metadata.page_count}</div>
                    <div class="metadata-item"><b>Author:</b> {extracted_content.metadata.author}</div>
                    <div class="metadata-item"><b>Title:</b> {extracted_content.metadata.title}</div>
                    <div class="metadata-item"><b>Creation Date:</b> {extracted_content.metadata.creation_date}</div>
                    <div class="metadata-item"><b>Last Modified:</b> {extracted_content.metadata.last_modified_date}</div>
                    <div class="metadata-item"><b>Language:</b> {extracted_content.metadata.language}</div>
                    <div class="metadata-item"><b>Extracted At:</b> {extracted_content.metadata.extracted_at}</div>
                </div>
            </div>

            <button class="collapsible">Full Text</button>
            <div class="content">
                <pre>{extracted_content.text}</pre>
            </div>

            <button class="collapsible">Sections and Segments</button>
            <div class="content">
        """

        for section in extracted_content.metadata.sections:
            html += f"""
                <div class="section">
                    <h3>{section.heading or 'Unnamed Section'}</h3>
                    <div class="metadata-item"><b>Section Index:</b> {section.section_index}</div>
                    <div class="metadata-item"><b>Character Count:</b> {section.char_count}</div>
                    <button class="collapsible">Segments ({len(section.segments)})</button>
                    <div class="content">
            """
            
            for segment in section.segments:
                html += f"""
                        <div class="segment">
                            <div class="metadata-item"><b>Type:</b> {segment.segment_type}</div>
                            <div class="metadata-item"><b>Index:</b> {segment.segment_index}</div>
                            <div class="metadata-item"><b>Characters:</b> {segment.char_count}</div>
                            <div class="metadata-item"><b>Metadata:</b> {segment.metadata}</div>
                            <pre>{segment.text}</pre>
                        </div>
                """
            
            html += """
                    </div>
                </div>
            """

        html += """
            </div>

            <script>
            var coll = document.getElementsByClassName("collapsible");
            var i;

            for (i = 0; i < coll.length; i++) {
                coll[i].addEventListener("click", function() {
                    this.classList.toggle("active");
                    var content = this.nextElementSibling;
                    if (content.style.display === "block") {
                        content.style.display = "none";
                    } else {
                        content.style.display = "block";
                    }
                });
            }
            </script>
        </body>
        </html>
        """

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

        return output_path 